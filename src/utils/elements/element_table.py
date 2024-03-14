import pathlib

import docx
import docx.oxml.simpletypes
from docx.shared import Inches, Pt, Cm


def create(*args):
    _columns: dict = args[0].get('columns', {})
    _h_columns: list[str] = _columns.get('columns_header', [])

    _table = docx.Document().add_table(rows=1, cols=len(_h_columns))

    for idx, _h_column in enumerate(_h_columns):
        width = modify_column(_table, idx, _columns[_h_column])

    return _table


def modify_column(table, column_idx: int, definition: dict):
    cell: docx.table.Cell = table.rows[0].cells[column_idx]
    ...
    cell.paragraphs[0].text = definition.get('label', '')

    alignment: object = get_alignment(definition=definition)
    if alignment:
        cell.paragraphs[0].alignment = alignment

    width: float = get_width(definition)

    if width:
        cell.width = Cm(width // 10)

    return width


def get_alignment(definition: dict) -> object | None:
    justify = definition.get('text-justify', None)
    if not justify:
        return None

    try:
        from utils.elements import justify_lookups
        justify = justify_lookups[justify]
        return justify
    except Exception:
        return None


def get_width(definition: dict) -> float | None:
    width = definition.get('width', None)

    if not width:
        return None

    width = width.split('m')[0]
    try:
        width_inches = float(width) / 25.4
        return width_inches
    except Exception:
        return None
